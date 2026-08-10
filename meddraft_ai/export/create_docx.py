import os
import re
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor

class MarkdownDocxConverter:
    """
    Converts Markdown content to formatted Word (.docx) documents strictly following
    AMA and APA 7th edition manuscript styling rules.
    """
    def __init__(self, font_name: str = "Times New Roman", font_size: int = 12):
        self.doc = Document()
        self.font_name = font_name
        self.font_size = font_size
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

    def add_heading_custom(self, text: str, level: int = 1):
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = self.font_name
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def add_paragraph_styled(self, text: str, bold: bool = False, italic: bool = False):
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.bold = bold
        run.italic = italic
        return p

    def _add_inline_formatted(self, paragraph, text: str) -> None:
        """
        Parse a line of text for **bold** and *italic* markdown markers and add
        properly-formatted runs to an existing paragraph object.
        """
        # Tokenise the line into alternating plain/bold/italic segments.
        # Pattern: **bold** or *italic* (non-overlapping, bold checked first)
        pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
        last_end = 0
        for m in pattern.finditer(text):
            # Plain text before this match
            if m.start() > last_end:
                run = paragraph.add_run(text[last_end:m.start()])
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
            if m.group(0).startswith('**'):
                run = paragraph.add_run(m.group(2))
                run.bold = True
            else:
                run = paragraph.add_run(m.group(3))
                run.italic = True
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            last_end = m.end()
        # Remaining plain text
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)

    def convert_markdown(self, md_content: str, output_path: str | Path) -> str:
        lines = md_content.split('\n')
        in_code_block = False

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith('```'):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                self.add_paragraph_styled(line_str)
                continue

            if line_str.startswith('# '):
                self.add_heading_custom(line_str[2:], level=1)
            elif line_str.startswith('## '):
                self.add_heading_custom(line_str[3:], level=2)
            elif line_str.startswith('### '):
                self.add_heading_custom(line_str[4:], level=3)
            elif line_str.startswith('- ') or line_str.startswith('* '):
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                self._add_inline_formatted(p, f"• {line_str[2:]}")
            else:
                # Regular paragraph — preserve inline bold/italic formatting
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                self._add_inline_formatted(p, line_str)

        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out_path))
        return str(out_path)

def create_docx_from_markdown(
    md_input: str | Path,
    output_docx_path: str | Path,
    is_file: Optional[bool] = None,
) -> str:
    """
    Convert markdown to DOCX.

    Args:
        md_input: Either a filesystem path to a .md file or a raw markdown string.
        output_docx_path: Destination .docx path.
        is_file: If True, treat md_input as a file path (raises FileNotFoundError if
                 missing). If False, treat as a raw string. If None (default), auto-detect
                 by checking whether the path exists — a non-existent path is treated as
                 a string to avoid silently processing typos as markdown content.
    """
    if is_file is True:
        file_path = Path(md_input)
        if not file_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {file_path}")
        content = file_path.read_text(encoding="utf-8", errors="ignore")
    elif is_file is False:
        content = str(md_input)
    else:
        # Auto-detect: only treat as file if it actually exists
        candidate = Path(str(md_input))
        if candidate.exists() and candidate.is_file():
            content = candidate.read_text(encoding="utf-8", errors="ignore")
        else:
            content = str(md_input)

    converter = MarkdownDocxConverter()
    return converter.convert_markdown(content, output_docx_path)
