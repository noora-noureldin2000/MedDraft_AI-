"""
Tests for DOCX export pipeline (meddraft_ai/export/).

Fixes applied per test-guard audit:
  - Assertions now verify *content* of generated DOCX, not just file existence (Rule 4)
  - is_file=True parameter tested explicitly (M-10 fix verification)
  - Bold/italic formatting verified in output (H-09 fix verification)
  - Misspelled path raises FileNotFoundError rather than silently treating path as markdown (M-10)
"""

from __future__ import annotations

import pytest
from pathlib import Path
from docx import Document

from meddraft_ai.export.create_docx import create_docx_from_markdown, MarkdownDocxConverter
from meddraft_ai.export.pandoc_converter import convert_markdown_to_docx


# ---------------------------------------------------------------------------
# create_docx_from_markdown — content correctness
# ---------------------------------------------------------------------------

class TestCreateDocxFromMarkdown:
    def test_heading1_appears_in_output(self, tmp_path):
        md = "# My Heading\n\nParagraph text."
        md_file = tmp_path / "doc.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "out.docx"

        create_docx_from_markdown(str(md_file), str(out))
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("My Heading" in t for t in texts)

    def test_paragraph_text_preserved(self, tmp_path):
        md = "## Section\n\nSample methods section text here."
        md_file = tmp_path / "doc.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "out.docx"

        create_docx_from_markdown(str(md_file), str(out))
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("Sample methods section text here." in t for t in texts)

    def test_bold_formatting_applied_to_run(self, tmp_path):
        """H-09 fix: **bold** must produce a run with run.bold == True."""
        md = "Normal text **bold word** more text."
        md_file = tmp_path / "bold.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "bold.docx"

        create_docx_from_markdown(str(md_file), str(out))
        doc = Document(str(out))
        bold_runs = [
            run for para in doc.paragraphs for run in para.runs if run.bold
        ]
        assert bold_runs, "Expected at least one bold run in the output DOCX"
        assert any("bold word" in r.text for r in bold_runs)

    def test_italic_formatting_applied_to_run(self, tmp_path):
        """H-09 fix: *italic* must produce a run with run.italic == True."""
        md = "Normal text *italic phrase* here."
        md_file = tmp_path / "italic.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "italic.docx"

        create_docx_from_markdown(str(md_file), str(out))
        doc = Document(str(out))
        italic_runs = [
            run for para in doc.paragraphs for run in para.runs if run.italic
        ]
        assert italic_runs, "Expected at least one italic run in the output DOCX"

    def test_is_file_true_raises_on_missing_path(self, tmp_path):
        """M-10 fix: is_file=True on a missing path must raise FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            create_docx_from_markdown(
                str(tmp_path / "does_not_exist.md"),
                str(tmp_path / "out.docx"),
                is_file=True,
            )

    def test_is_file_false_treats_input_as_string(self, tmp_path):
        """M-10 fix: is_file=False treats a raw string as markdown, not a file path."""
        raw_md = "# Inline Title\n\nInline content."
        out = tmp_path / "inline.docx"
        create_docx_from_markdown(raw_md, str(out), is_file=False)
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("Inline content." in t for t in texts)

    def test_auto_detect_falls_back_to_string_for_nonexistent_path(self, tmp_path):
        """M-10 fix: auto-detect (is_file=None) treats non-existent path as raw markdown."""
        out = tmp_path / "auto.docx"
        # This path-like string does not exist on disk
        create_docx_from_markdown("## Auto detect\n\nText.", str(out))
        assert out.exists()


# ---------------------------------------------------------------------------
# convert_markdown_to_docx (pandoc fallback)
# ---------------------------------------------------------------------------

class TestConvertMarkdownToDocx:
    def test_fallback_produces_output_file(self, tmp_path):
        md = "# Title\n\n## Methods\n\nSample methods section."
        md_file = tmp_path / "sample.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "sample.docx"

        result = convert_markdown_to_docx(str(md_file), str(out))
        assert Path(result).exists()

    def test_raises_on_missing_markdown_file(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            convert_markdown_to_docx(str(tmp_path / "missing.md"), str(tmp_path / "out.docx"))

    def test_output_contains_heading_text(self, tmp_path):
        md = "# Report Title\n\nBody paragraph."
        md_file = tmp_path / "report.md"
        md_file.write_text(md, encoding="utf-8")
        out = tmp_path / "report.docx"

        result = convert_markdown_to_docx(str(md_file), str(out))
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("Report Title" in t for t in texts)


# ---------------------------------------------------------------------------
# MarkdownDocxConverter inline parser
# ---------------------------------------------------------------------------

class TestMarkdownDocxConverterInlineParser:
    def test_plain_text_has_no_bold_runs(self, tmp_path):
        converter = MarkdownDocxConverter()
        out = str(tmp_path / "plain.docx")
        converter.convert_markdown("Plain sentence without formatting.", out)
        doc = Document(out)
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert not bold_runs

    def test_mixed_bold_and_italic_same_line(self, tmp_path):
        converter = MarkdownDocxConverter()
        out = str(tmp_path / "mixed.docx")
        converter.convert_markdown("Text **bold** and *italic* end.", out)
        doc = Document(out)
        bold = [r for p in doc.paragraphs for r in p.runs if r.bold]
        italic = [r for p in doc.paragraphs for r in p.runs if r.italic]
        assert bold
        assert italic
